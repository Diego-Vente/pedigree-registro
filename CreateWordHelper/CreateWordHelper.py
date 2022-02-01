import docx
import os
import sys
from os.path import expanduser
from database.database import WorkspaceData
import json
import subprocess
import os.path
from os import path
import shutil
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from Models.models import *


class CreateWordHelper:
    def __init__(self):
        # Chequear y generar Carpetas structure

        self.home = expanduser("~")
        print("Home")
        print(self.home)
        self.main_app_path = self.home + "\\Registros Pedigree App"
        self.pedigree_app_path = self.main_app_path + "\\Pedigrees"
        self.clientes_app_path = self.main_app_path + "\\Afijos"
        self.backup_app_path = self.main_app_path + "\\Backups"

        if not path.exists(self.main_app_path):
            os.mkdir(self.main_app_path)  # GENERO CARPETA MAIN PARA ARCHIVOS
        if not path.exists(self.pedigree_app_path):
            os.mkdir(self.pedigree_app_path)
        if not path.exists(self.clientes_app_path):
            os.mkdir(self.clientes_app_path)
        if not path.exists(self.backup_app_path):
            os.mkdir(self.backup_app_path)

        # Pedigree base Object
        self.pedigree_base_doc = docx.Document('pedigree_base_anverso.docx')
        self.pedigree_base_reverso_doc = docx.Document('pedigree_base_reverso.docx')
        # Afijo cliente bas Object
        self.afijo_cliente_base_doc = docx.Document('afijo_cliente_base.docx')

        # sections = self.pedigree_base_doc.sections
        # section = sections[0]  # Obtener primera seccion (Pagina)
        # print("--- tables ----")
        # tables = self.pedigree_base_doc.tables
        # print(tables)
        # print(tables[0].cell)
        # print(len(tables))
        #
        # # Afijo base Object
        # sections1 = self.pedigree_base_reverso_doc.sections
        # section = sections1[0]  # Obtener primera seccion (Pagina)
        # print("--- tables ----")
        # tables1 = self.pedigree_base_doc.tables
        # print(tables1)
        # print(tables1[0].cell)
        # print(len(tables1))
        #
        # for table in tables1:
        #     print(table.cell(0, 0).text)

        # Root carpetas
        # Pedigrees
        #   - NOMBRE PERRO AFIJO -> Pedigree, Tarjeta de identificacion, (posiblemente screenshot).
        # Afijos
        #   - NOMBRE CLIENTE AFIJO -> Afijo cliente, Tarjeta de identificacion

    def createPedigree(self, pedigree_data: PedigreeData):
        test_string = "mama/color de mama"
        result = test_string.split("/")
        print(result)
        print(len(result))

        # Obtener data y generar archivos
        new_line = "\n"
        cachorro_name = pedigree_data.nombre_cachorro
        afijo_cachorro = str(pedigree_data.certificado_code)
        carpeta_name = cachorro_name.upper() + " - " + afijo_cachorro.upper()
        path_file = self.pedigree_app_path + f'\\{carpeta_name}'

        genealogia = pedigree_data.pedigree_gene_data
        gene_madre = genealogia["madre"]
        gene_padre = genealogia["padre"]
        # Create carpeta for this Pedigree if not exists
        if not path.exists(path_file):
            os.mkdir(path_file)

        # 1 -> Registro NRO
        registro_nro = self.setAfijoFormat(int(pedigree_data.certificado_code), "nro_code")
        # 3 -> nombre,codigo,raza
        tabla_nombre = cachorro_name.upper() + new_line + new_line\
                       + self.setAfijoFormat(int(pedigree_data.certificado_code), "name") + new_line + new_line\
                       + pedigree_data.raza.upper()
        # 5 -> chip, criador, propietario
        tabla_detalles_cliente = pedigree_data.chip_code + new_line + new_line +\
                                 pedigree_data.criador.upper() + new_line + new_line +\
                                 pedigree_data.propietario.upper()
        # 17 -> sexo, color, f_naciemiento
        tabla_detalles_sexo = pedigree_data.sexo.upper() + new_line + new_line +\
                              pedigree_data.color.upper() + new_line + new_line +\
                              pedigree_data.nacimiento

        # 2 -> madre/m
        # 4 -> padre/p
        # 6 -> madre / am
        # 7 -> madre / ap
        # 8 -> padre / am
        # 9 -> padre / ap
        # 10 -> madre / bm_p
        # 11 -> madre / bp_m
        # 12 -> padre / bp_m
        # 13 -> padre / bp_p
        # 14 -> madre / bm_m
        # 15 -> padre / bm_p
        # 16 -> padre / bm_m


        # 19 -> madre / bp_p

        # 18 - 1 -> madre / tm_bm_m
        # 18 - 2 -> madre / tm_bm_p
        # 18 - 3 -> madre / tm_bp_m
        # 18 - 4 -> madre / tm_bp_p
        # 18 - 5 -> madre / tp_bm_m
        # 18 - 6 -> madre / tp_bm_p
        # 18 - 7 -> madre / tp_bp_m
        # 18 - 8 -> madre / tp_bp_p

        # 18 - 9 -> padre / tm_bm_m
        # 18 - 10 -> padre / tm_bm_p
        # 18 - 11 -> padre / tm_bp_m
        # 18 - 12 -> padre / tm_bp_p
        # 18 - 13 -> padre / tp_bm_m
        # 18 - 14 -> padre / tp_bm_p
        # 18 - 15 -> padre / tp_bp_m
        # 18 - 16 -> padre / tp_bp_p

        # Create Pedigree document
        # Reverso
        sections = self.pedigree_base_doc.sections
        section = sections[0]  # Obtener primera seccion (Pagina)
        print("--- tables ----")
        tables = self.pedigree_base_doc.tables
        print(tables)
        print(tables[0].cell)
        print(len(tables))

        current_cell = ""
        counter = 1
        for table in tables:
            counter_a = 1
            for row in table.rows:
                for cell in row.cells:
                    current_cell = str(counter) + " - " + str(counter_a)
#                    cell.text = current_cell
                    font_size = 8
                    text_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    if current_cell == "1 - 1":
                        cell.text = registro_nro
                        font_size = 16
                        text_alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif current_cell == "4 - 1":
                        cell.text = tabla_nombre
                        text_alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif current_cell == "2 - 1":
                        cell.text = tabla_detalles_sexo
                        text_alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif current_cell == "3 - 1":
                        cell.text = tabla_detalles_cliente
                        text_alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                    elif current_cell == "9 - 1":
                        cell.text = self.cleanNameColor(gene_madre["m"])
                    elif current_cell == "19 - 1":
                        cell.text = self.cleanNameColor(gene_padre["p"])
                    elif current_cell == "8 - 1":
                        cell.text = self.cleanNameColor(gene_madre["am"])
                    elif current_cell == "12 - 1":
                        cell.text = self.cleanNameColor(gene_madre["ap"])
                    elif current_cell == "17 - 1":
                        cell.text = self.cleanNameColor(gene_padre["am"])
                    elif current_cell == "18 - 1":
                        cell.text = self.cleanNameColor(gene_padre["ap"])

                    elif current_cell == "6 - 1":
                        cell.text = self.cleanNameColor(gene_madre["bm_m"])
                    elif current_cell == "7 - 1":
                        cell.text = self.cleanNameColor(gene_madre["bm_p"])
                    elif current_cell == "10 - 1":
                        cell.text = self.cleanNameColor(gene_madre["bp_m"])
                    elif current_cell == "11 - 1":
                        cell.text = self.cleanNameColor(gene_madre["bp_p"])
                    elif current_cell == "13 - 1":
                        cell.text = self.cleanNameColor(gene_padre["bm_m"])
                    elif current_cell == "14 - 1":
                        cell.text = self.cleanNameColor(gene_padre["bm_p"])
                    elif current_cell == "15 - 1":
                        cell.text = self.cleanNameColor(gene_padre["bp_m"])
                    elif current_cell == "16 - 1":
                        cell.text = self.cleanNameColor(gene_padre["bp_p"])

                    elif current_cell == "5 - 1":
                        cell.text = self.cleanNameColor(gene_madre["tm_bm_m"])
                    elif current_cell == "5 - 2":
                        cell.text = self.cleanNameColor(gene_madre["tm_bm_p"])
                    elif current_cell == "5 - 3":
                        cell.text = self.cleanNameColor(gene_madre["tm_bp_m"])
                    elif current_cell == "5 - 4":
                        cell.text = self.cleanNameColor(gene_madre["tm_bp_p"])
                    elif current_cell == "5 - 5":
                        cell.text = self.cleanNameColor(gene_madre["tp_bm_m"])
                    elif current_cell == "5 - 6":
                        cell.text = self.cleanNameColor(gene_madre["tp_bm_p"])
                    elif current_cell == "5 - 7":
                        cell.text = self.cleanNameColor(gene_madre["tp_bp_m"])
                    elif current_cell == "5 - 8":
                        cell.text = self.cleanNameColor(gene_madre["tp_bp_p"])
                    elif current_cell == "5 - 9":
                        cell.text = self.cleanNameColor(gene_padre["tm_bm_m"])
                    elif current_cell == "5 - 10":
                        cell.text = self.cleanNameColor(gene_padre["tm_bm_p"])
                    elif current_cell == "5 - 11":
                        cell.text = self.cleanNameColor(gene_padre["tm_bp_m"])
                    elif current_cell == "5 - 12":
                        cell.text = self.cleanNameColor(gene_padre["tm_bp_p"])
                    elif current_cell == "5 - 13":
                        cell.text = self.cleanNameColor(gene_padre["tp_bm_m"])
                    elif current_cell == "5 - 14":
                        cell.text = self.cleanNameColor(gene_padre["tp_bm_p"])
                    elif current_cell == "5 - 15":
                        cell.text = self.cleanNameColor(gene_padre["tp_bp_m"])
                    elif current_cell == "5 - 16":
                        cell.text = self.cleanNameColor(gene_padre["tp_bp_p"])
                    else:
                        cell.text = "Error"
                        print(cell.text)

                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        paragraph.alignment = text_alignment
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(font_size)
                            font.name = 'Arial'
                counter_a += 1
            counter += 1

        # Check if path exits.
        pedigree_doc_path = path_file+"\\registro_pedigree_reverso.docx"

        # Anverso
        sections_reverso = self.pedigree_base_reverso_doc.sections
        section = sections_reverso[0]  # Obtener primera seccion (Pagina)
        print("--- tables ----")
        tables_reverso = self.pedigree_base_reverso_doc.tables
        pedigree_anverso_doc_path = path_file + "\\registro_pedigree_anverso.docx"
        for table in tables_reverso:
            print(table.cell(0, 0).text)

            table.cell(0, 0).text = pedigree_data.nombre_cachorro.upper() + new_line + new_line +\
                                    str(pedigree_data.certificado_code) + new_line + new_line + \
                                    pedigree_data.raza.upper() + new_line + new_line + \
                                    pedigree_data.sexo.upper() + new_line + new_line + \
                                    pedigree_data.color.upper() + new_line + new_line + \
                                    pedigree_data.nacimiento + new_line + new_line + \
                                    self.getNameCachorro(pedigree_data.nombre_madre) + new_line + new_line + \
                                    pedigree_data.afijo_madre + new_line + new_line + \
                                    self.getNameCachorro(pedigree_data.nombre_padre.upper()) + new_line + new_line + \
                                    pedigree_data.afijo_padre + new_line + new_line + \
                                    pedigree_data.chip_code + new_line + new_line + \
                                    pedigree_data.criador.upper() + new_line + new_line + \
                                    pedigree_data.propietario.upper()
            paragraphs = table.cell(0, 0).paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)
                    font.name = 'Arial'

        self.pedigree_base_reverso_doc.save(pedigree_anverso_doc_path)  # This overwrite by defect.
        self.pedigree_base_doc.save(pedigree_doc_path)  # This overwrite by defect.

        return pedigree_doc_path

    def createClienteAfijo(self, clienteData):
        print("----------------------CLIENTE---------------------------")
        # Afijo base Object

        cliente_name = clienteData[0]
        cliente_criador_name = clienteData[1]
        cliente_afijo = clienteData[2]

        carpeta_name = cliente_name.upper() + " - " + cliente_afijo.replace('/', '')
        print("carpeta name:")
        print(carpeta_name)
        path_file = self.clientes_app_path + f'\\{carpeta_name}'
        afijo_cliente_doc_path = path_file + "\\afijo.docx"

        new_line = "\n"
        if not path.exists(path_file):
            os.mkdir(path_file)

        tables = self.afijo_cliente_base_doc.tables

        # 1 -> small Details
        # 2 -> big details

        details_afijo = cliente_name.upper() + new_line + new_line + \
                        cliente_criador_name.upper() + new_line + new_line +\
                        cliente_afijo

        counter = 1
        for table in tables:
            if counter == 1:
                table.cell(0, 0).text = details_afijo
                paragraphs = table.cell(0, 0).paragraphs
                font_size = 7
                for paragraph in paragraphs:
                    print("one paraph")
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for run in paragraph.runs:
                        print("two paraph")
                        font = run.font
                        font.size = Pt(font_size)
                        font.name = 'Arial'
            elif counter == 2:
                table.cell(0, 0).text = details_afijo
                font_size = 22
                paragraphs = table.cell(0, 0).paragraphs
                for paragraph in paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(font_size)
                        font.name = 'Arial'

            counter += 1

        self.afijo_cliente_base_doc.save(afijo_cliente_doc_path)  # This overwrite by defect.
        return afijo_cliente_doc_path

    def deleteFile(self, path):
        shutil.rmtree(path)

    def cleanNameColor(self, name_color):
        new_line = "\n"
        text_list = name_color.split("/")
        if len(text_list) > 1:
            name = text_list[0].upper()
            color = text_list[1].upper()
            text = name + new_line + color
            return text
        else:
            return name_color

    def getNameCachorro(self, name_color):
        text_list = name_color.split("/")
        if len(text_list) > 1:
            name = text_list[0].upper()
            return name
        else:
            return name_color

    def setAfijoFormat(self, afijo, type):
        if type == "name":
            # Consideramos 7 ceros
            if afijo < 100000:
                afijo = "#0000000" + str(afijo) + "-PER"
            elif afijo < 1000000:
                afijo = "#000000" + str(afijo) + "-PER"
            elif afijo < 10000000:
                afijo = "#00000" + str(afijo) + "-PER"
            elif afijo < 100000000:
                afijo = "#0000" + str(afijo) + "-PER"
            elif afijo < 1000000000:
                afijo = "#000" + str(afijo) + "-PER"
            else:
                afijo = "#0000000" + str(afijo) + "-PER"
        elif type == "nro_code":
            afijo = "PER-" + str(afijo)
        return afijo
